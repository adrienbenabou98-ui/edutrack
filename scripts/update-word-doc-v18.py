import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

path = r'C:\Users\adrie\Documents\EduTrack-Credentials.docx'
doc = Document(path)

for para in doc.paragraphs:
    if 'Last updated:' in para.text and ('v1.7.0' in para.text or 'v1.6.0' in para.text or 'v1.5.0' in para.text):
        for run in para.runs:
            if 'Last updated:' in run.text or 'v1.' in run.text:
                run.text = 'Last updated: 2026-05-26 (v1.8.0 - 10 CRM+Canva features)'
                break
        break

for para in doc.paragraphs:
    if 'v1.7.0' in para.text and 'releases/tag' in para.text:
        for run in para.runs:
            if 'v1.7.0' in run.text:
                run.text = run.text.replace('v1.7.0', 'v1.8.0')
    elif 'v1.6.0' in para.text and 'releases/tag' in para.text:
        for run in para.runs:
            if 'v1.6.0' in run.text:
                run.text = run.text.replace('v1.6.0', 'v1.8.0')

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if 'v1.7.0' in run.text and 'releases/tag' in run.text:
                        run.text = run.text.replace('v1.7.0', 'v1.8.0')

# Append new features section at end
doc.add_heading('v1.8.0 New Features (2026-05-26)', level=2)
doc.add_paragraph('Assignment Template Library - /teacher/templates - save and reuse assignment structures.')
doc.add_paragraph('Lesson Planner - /teacher/planner - weekly 5-day x 6-period grid.')
doc.add_paragraph('Intervention Pipeline - /teacher/interventions - Kanban board (Monitoring/Contacted/Improving/Resolved).')
doc.add_paragraph('Announcement Designer - /teacher/announcements - write and send announcements with templates.')
doc.add_paragraph('Student Timeline - new tab on StudentGradeDetail - per-student activity log.')
doc.add_paragraph('Parent Communication Log - new tab on StudentGradeDetail - log parent contacts.')
doc.add_paragraph('Engagement Scores - HIGH/MEDIUM/LOW badges in ClassroomDetail > Students tab.')
doc.add_paragraph('At-Risk Alerts - widget on teacher Dashboard (2+ missing or avg < 50%).')
doc.add_paragraph('Automated Reminders - daily scheduler sends notifications for assignments due within 48h.')
doc.add_paragraph('New DB models: AssignmentTemplate, TemplateQuestion, LessonPlan, ParentContact, Intervention, TeacherAnnouncement.')
doc.add_paragraph('All new endpoints: authenticated + requireRole(TEACHER) + ownership check. No PII leaks.')

doc.save(path)
print('Done')
